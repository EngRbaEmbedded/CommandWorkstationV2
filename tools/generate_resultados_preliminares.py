#!/usr/bin/env python3
"""Generate MBA Resultados Preliminares docx from official USP/Esalq template."""

from __future__ import annotations

import json
import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

PROJECT_ROOT = Path(__file__).resolve().parents[1]
RESULTS_JSON = PROJECT_ROOT / "docs" / "test_results.json"
TEMPLATE = Path("/home/rbabynho/MBA/1 Resultados Preliminares_09-06--16-06.docx")
OUTPUT = Path("/home/rbabynho/MBA/Resultados Preliminares - Rodrigo Borraschi Antonio.docx")

COURSE = "MBA em Engenharia de Software"
DEFENSE_YEAR = "2026"
TITLE = "Command Workstation: painel unificado de comandos para engenharia de software"
AUTHORS = "Rodrigo Borraschi Antonio¹*; Arthur Pinheiro De Araújo Costa²"
FOOTNOTE1 = (
    "¹* MBA em Engenharia de Software — USP/Esalq. "
    "autor correspondente: rodrigo.borraschi@email.com"
)
FOOTNOTE2 = (
    "² Doutor. MBA em Engenharia de Software — USP/Esalq. "
    "e-mail: orientador@email.com"
)
HEADER = (
    f"Trabalho de Conclusão de Curso apresentado para obtenção do título de "
    f"especialista em {COURSE} - {DEFENSE_YEAR}"
)

RESUMO = (
    "Este trabalho apresentou o desenvolvimento preliminar do Command Workstation, "
    "protótipo desktop concebido como painel unificado (Single Pane of Glass) para "
    "reduzir a alternância entre ferramentas no fluxo de engenharia de software. "
    "Após reorientação do escopo, com anuência do orientador, concentrou-se um núcleo "
    "mínimo funcional composto por execução local de comandos, terminal integrado, "
    "operações Git básicas e persistência em SQLite. A validação técnica compreendeu "
    "roteiro de nove testes funcionais, todos aprovados, e análise do histórico de "
    "execuções registrado automaticamente no banco SQLite. Concluiu-se pela viabilidade "
    "técnica do núcleo mínimo proposto nesta etapa preliminar."
)
PALAVRAS_CHAVE = (
    "painel unificado; protótipo de software; carga cognitiva; "
    "engenharia de software; SQLite."
)

CONSIDERACOES_FINAIS = (
    "Até esta etapa, concluiu-se que a reconstrução enxuta do Command Workstation "
    "atendeu ao núcleo mínimo definido com o orientador, demonstrando viabilidade "
    "técnica para execução de comandos, integração Git básica, terminal integrado e "
    "persistência local. A avaliação de impacto sobre produtividade e carga cognitiva "
    "permanece pendente para o TCC final, mediante testes de usabilidade com participantes "
    "voluntários."
)


def set_run_font(run, *, size: int = 11, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rfonts.set(qn("w:eastAsia"), "Arial")


def clear_paragraph(paragraph: Paragraph) -> None:
    paragraph.text = ""


def write_paragraph(
    paragraph: Paragraph,
    text: str,
    *,
    bold: bool = False,
    size: int = 11,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.paragraph_format.alignment = align


def format_body(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        set_run_font(run)


def format_subheading(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, bold=True)


def format_section_heading(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run, bold=True)


def format_resumo_body(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        set_run_font(run)


def format_footnote(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run, size=9)


def format_reference(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run)


def format_table_caption(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        set_run_font(run)


def format_table_source(paragraph: Paragraph) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = "Normal"
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_body_paragraphs_after(anchor: Paragraph, texts: list[str]) -> Paragraph:
    current = anchor
    for text in texts:
        current = insert_paragraph_after(current, text)
        format_body(current)
    return current


def insert_subheading_after(anchor: Paragraph, text: str) -> Paragraph:
    para = insert_paragraph_after(anchor, text)
    format_subheading(para)
    return para


def _set_cell_text(cell, text: str, *, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = str(text)
    for p in cell.paragraphs:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.alignment = alignment
        for run in p.runs:
            set_run_font(run, bold=False)


def _apply_minimal_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "right", "insideV", "insideH"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)

    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(borders)

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_borders = OxmlElement("w:tcBorders")
            for edge in ("left", "right", "insideV"):
                element = OxmlElement(f"w:{edge}")
                element.set(qn("w:val"), "nil")
                cell_borders.append(element)
            if row_idx == 0:
                top = OxmlElement("w:top")
                top.set(qn("w:val"), "single")
                top.set(qn("w:sz"), "8")
                top.set(qn("w:color"), "000000")
                cell_borders.append(top)
            else:
                top = OxmlElement("w:top")
                top.set(qn("w:val"), "nil")
                cell_borders.append(top)
            bottom = OxmlElement("w:bottom")
            if row_idx == len(table.rows) - 1:
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "8")
                bottom.set(qn("w:color"), "000000")
            else:
                bottom.set(qn("w:val"), "nil")
            cell_borders.append(bottom)
            existing = tc_pr.find(qn("w:tcBorders"))
            if existing is not None:
                tc_pr.remove(existing)
            tc_pr.append(cell_borders)


def insert_table_with_caption(
    anchor: Paragraph,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
    *,
    numeric_cols: set[int] | None = None,
) -> Paragraph:
    numeric_cols = numeric_cols or set()
    cap = insert_paragraph_after(anchor, caption)
    format_table_caption(cap)

    doc = anchor.part.document
    table = doc.add_table(rows=1, cols=len(headers))
    _apply_minimal_table_borders(table)

    for i, header in enumerate(headers):
        if i == 0:
            align = WD_ALIGN_PARAGRAPH.LEFT
        elif i in numeric_cols:
            align = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            align = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_text(table.rows[0].cells[i], header, alignment=align)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if i == 0:
                align = WD_ALIGN_PARAGRAPH.LEFT
            elif i in numeric_cols:
                align = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                align = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_cell_text(cells[i], value, alignment=align)

    cap._p.addnext(table._tbl)

    fonte_el = OxmlElement("w:p")
    table._tbl.addnext(fonte_el)
    fonte = Paragraph(fonte_el, anchor._parent)
    fonte.add_run("Fonte: Resultados originais da pesquisa")
    format_table_source(fonte)

    spacer = OxmlElement("w:p")
    fonte_el.addnext(spacer)
    return Paragraph(spacer, anchor._parent)


def find_paragraph(doc: Document, starts_with: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(starts_with):
            return paragraph
    raise ValueError(f"Paragraph not found: {starts_with!r}")


def remove_instruction_paragraphs(doc: Document) -> None:
    prefixes = (
        "Atenção:",
        "O título da seção",
        "A introdução deve ser redigida",
        "Os TCCs dos cursos",
        "As pesquisas envolvendo",
        "Neste tópico deve ser apresentada",
        "O(s) nome(s) do(s) participante(s)",
        "Tópico obrigatório para o depósito",
        "Neste tópico deverão ser listadas",
        "Apêndices são textos",
        "O TCC deverá conter no máximo 30",
        "O título da seção Resultados Preliminares deve",
        "O título da seção Agradecimentos",
    )
    to_remove = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if any(text.startswith(prefix) for prefix in prefixes):
            to_remove.append(paragraph)
    for paragraph in to_remove:
        delete_paragraph(paragraph)


def remove_optional_template_sections(doc: Document) -> None:
    optional_starts = (
        "Conclusão(ões) ou Considerações Finais",
        "Agradecimento (opcional, 1 parágrafo, bem sucinto)",
        "Apêndice ou Anexo (opcional)",
    )
    collecting = False
    to_remove = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if any(text.startswith(prefix) for prefix in optional_starts):
            collecting = True
        if text.startswith("Referências"):
            collecting = False
        if collecting:
            to_remove.append(paragraph)
    for paragraph in to_remove:
        delete_paragraph(paragraph)


def setup_folha_de_rosto(doc: Document) -> None:
    write_paragraph(doc.paragraphs[0], HEADER, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_run_font(doc.paragraphs[0].runs[0], size=11)

    title_para = doc.paragraphs[1] if len(doc.paragraphs) > 1 else insert_paragraph_after(doc.paragraphs[0])
    write_paragraph(title_para, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(title_para.runs[0], size=11, bold=True)

    authors_para = find_paragraph(doc, "nome completo aluno")
    footnote1_para = find_paragraph(doc, "1* Titulação")
    footnote2_para = find_paragraph(doc, "2 Titulação")

    write_paragraph(authors_para, AUTHORS, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(authors_para.runs[0], size=11)

    write_paragraph(footnote1_para, FOOTNOTE1)
    format_footnote(footnote1_para)

    write_paragraph(footnote2_para, FOOTNOTE2)
    format_footnote(footnote2_para)


def setup_second_page(doc: Document) -> None:
    title2 = find_paragraph(doc, "Título do trabalho de conclusão de curso")
    write_paragraph(title2, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(title2.runs[0], size=11, bold=True)

    resumo_heading = find_paragraph(doc, "Resumo (ou Sumário Executivo)")
    write_paragraph(resumo_heading, "Resumo", bold=True)
    format_section_heading(resumo_heading)

    resumo_body = find_paragraph(doc, "Tópico obrigatório para o depósito")
    write_paragraph(resumo_body, RESUMO)
    format_resumo_body(resumo_body)

    kw_para = find_paragraph(doc, "Palavras-chave:")
    clear_paragraph(kw_para)
    label = kw_para.add_run("Palavras-chave: ")
    set_run_font(label, bold=True)
    words = kw_para.add_run(PALAVRAS_CHAVE)
    set_run_font(words)
    format_resumo_body(kw_para)


def build_content(data: dict) -> dict:
    metrics = data["metrics"]
    summary = data["summary"]
    tests = data["test_cases"]
    today = date.today().strftime("%d/%m/%Y")
    exit_codes = metrics.get("por_exit_code", {})
    exit_desc = ", ".join(f"exit {k}: {v}" for k, v in exit_codes.items())
    tempo = metrics.get("tempo_medio_ms_sucesso", 0)

    intro = [
        "No cenário contemporâneo de desenvolvimento de software, profissionais alternaram "
        "frequentemente entre terminais, sistemas de controle de versão, ferramentas de build "
        "e outras interfaces especializadas. Essa fragmentação impôs custos atencionais "
        "significativos: após cada troca de contexto, o indivíduo necessitou de tempo adicional "
        "para retomar o estado cognitivo anterior, o que elevou o estresse e reduziu a agilidade "
        "na execução de tarefas complexas (Mark; Gudith; Klocke, 2008). Revisões recentes "
        "associaram a dispersão de informações entre múltiplas interfaces a maior carga cognitiva "
        "e menor compreensão global dos processos, quando comparadas a abordagens integradas "
        "(Schreiber; Abbad-Andaloussi; Weber, 2023; Shanmugasundaram; Tamilarasu, 2023). "
        "A teoria da carga cognitiva (Sweller, 1988) sustentou que interfaces mal projetadas "
        "consumiram recursos da memória de trabalho que deveriam ser direcionados à resolução "
        "do problema em si.",
        "Diante desse cenário, propôs-se o Command Workstation — aplicação desktop que "
        "implementou o conceito de Single Pane of Glass, reunindo em uma única interface o "
        "acesso a comandos locais, terminal integrado, operações Git e persistência de "
        "configurações e histórico de execuções. O trabalho foi inicialmente delineado para "
        "desenvolvimento e avaliação em ambiente corporativo com equipes de engenharia de "
        "software. Em razão da perda de acesso ao ambiente e aos artefatos originais, o "
        "escopo foi reorientado, com anuência do orientador, para a construção e validação "
        "técnica de um protótipo funcional (prova de conceito), preservando o tema central "
        "do painel unificado de comandos.",
        "O objetivo deste trabalho consistiu em desenvolver o protótipo Command Workstation e "
        "validar tecnicamente seu núcleo mínimo — painel unificado, execução local de comandos, "
        "terminal embutido, integração básica com Git e persistência em SQLite — como "
        "alternativa viável à alternância entre múltiplas ferramentas no fluxo de trabalho "
        "de desenvolvimento de software.",
    ]

    return {
        "intro": intro,
        "metodologia_intro": [
            "A pesquisa manteve caráter aplicado, porém o delineamento foi ajustado de estudo "
            "de caso corporativo para desenvolvimento iterativo de protótipo e validação técnica "
            "funcional. O trabalho dividiu-se em duas frentes complementares: (a) reconstrução "
            "do software e (b) verificação do comportamento do núcleo mínimo da proposta.",
        ],
        "metodologia_dev": [
            "O protótipo foi reconstruído como aplicação desktop multiplataforma utilizando C++17, "
            "Qt 6 (Quick/QML) para a interface gráfica, CMake para build e SQLite (módulo Qt SQL) "
            "para persistência local. A arquitetura separou camada de apresentação (QML), "
            "orquestração (WorkstationBridge), execução de comandos (CommandRunner), integração "
            "com terminal do sistema (TerminalLauncher) e persistência (DatabaseManager). "
            "A execução de comandos utilizou subprocessos do sistema operacional (QProcess), "
            "com saída padrão e de erro exibidas em tempo real no terminal embutido.",
            "O núcleo mínimo implementado compreendeu: seleção de pasta de trabalho; presets de "
            "comandos agrupados por categorias (General, Build/check, Quality) e presets "
            "personalizados editáveis pelo usuário; terminal embutido e opção de terminal "
            "interativo do sistema; integração Git (branch, status, pull, push, commit e "
            "descarte de alterações); e registro automático de execuções na tabela run_log "
            "(comando, diretório, código de saída, duração e trecho da saída). Funcionalidades "
            "do escopo original não replicadas nesta fase incluíram comandos corporativos "
            "específicos de simulação, workflow de oito etapas, integração com sistemas internos "
            "de gestão e avaliação junto a equipes no ambiente empresarial original.",
        ],
        "metodologia_val": [
            f"A validação preliminar concentrou-se em roteiro estruturado de testes funcionais "
            f"(TF-01 a TF-09), executado em {today}, verificando inicialização e conexão SQLite, "
            "detecção de repositório Git, execução de comandos tipo preset, persistência de presets "
            "personalizados, comandos customizados, interrupção de processos, disponibilidade do "
            "binário compilado e ambiente de build. Complementarmente, foram analisados registros "
            "históricos persistidos automaticamente no banco SQLite (tabelas settings, presets e run_log).",
        ],
        "resultados_dev": [
            "Até o momento, obteve-se um protótipo executável do Command Workstation, compilado "
            "para Linux (binário appCommand_Workstation), com interface unificada organizada "
            "em painel lateral de presets e histórico, área de informações Git e terminal "
            "integrado. O banco SQLite local armazenou configurações, presets e histórico de execuções.",
            f"Foram identificadas {metrics['total_presets']} entradas de presets (padrão e "
            f"personalizados) e {metrics['total_execucoes']} registros no histórico de execuções "
            "(run_log). A Tabela 1 resume o status de implementação dos componentes previstos para o MVP.",
        ],
        "resultados_val": [
            f"Executaram-se {summary['total']} casos de teste funcionais, dos quais "
            f"{summary['passed']} foram aprovados e {summary['failed']} reprovados. "
            "A Tabela 2 consolida os resultados obtidos.",
            f"Na análise documental do histórico de execuções, observou-se distribuição por "
            f"código de saída ({exit_desc}). O tempo médio de execução bem-sucedida foi de "
            f"{tempo} ms (comandos com exit code 0). A Tabela 3 apresenta os comandos mais "
            "registrados no histórico.",
            "Entre os resultados destacados, verificou-se: (i) conexão estável ao SQLite com "
            "schema esperado; (ii) detecção de repositório Git no diretório de trabalho "
            "selecionado; (iii) execução assíncrona de comandos com registro automático de "
            "duração e código de saída; (iv) persistência de presets personalizados mediante "
            "operações INSERT/SELECT no banco; e (v) possibilidade de interrupção de processos "
            "em execução (equivalente ao botão Stop da interface).",
            "Limitações dos resultados preliminares: os achados referiram-se à viabilidade "
            "técnica do núcleo mínimo da proposta. Não foi possível conduzir a avaliação "
            "pré e pós-intervenção junto a equipes corporativas prevista no projeto original. "
            "A dimensão de impacto sobre carga cognitiva e produtividade percebida permaneceu "
            "como perspectiva para o TCC final, conforme delimitação acordada com o orientador.",
        ],
        "refs": [
            "AMERICAN PSYCHOLOGICAL ASSOCIATION. Multitasking: switching costs. APA, 2026. "
            "Disponível em: https://www.apa.org/topics/research/multitasking. Acesso em: 05 mar. 2026.",
            "MARK, G.; GUDITH, D.; KLOCKE, U. The cost of interrupted work: more speed and stress. "
            "In: CHI Conference on Human Factors in Computing Systems, 2008, Florence. Anais... p. 107-110.",
            "SCHREIBER, C.; ABBAD-ANDALOUSSI, A.; WEBER, B. On the cognitive effects of abstraction "
            "and fragmentation in modularized process models. In: International Conference on Business "
            "Process Management, 2023, Utrecht. Anais... p. 359-376.",
            "SHANMUGASUNDARAM, M. S. M.; TAMILARASU, A. T. The impact of digital technology, social "
            "media, and artificial intelligence on cognitive functions: a review. Frontiers in Cognition, "
            "2023.",
            "SOMMERVILLE, I. Software Engineering. 10. ed. Pearson, 2016.",
            "SWELLER, J. Cognitive load during problem solving: effects on learning. Cognitive Science, "
            "v. 12, n. 2, p. 257-285, 1988.",
        ],
        "table1": [
            ["Painel unificado", "Implementado", "Interface QML com presets, Git e terminal"],
            ["Presets padrão", "Implementado", "Categorias General, Build/check e Quality"],
            ["Presets personalizados", "Implementado", "CRUD na categoria Personalizados"],
            ["Terminal embutido", "Implementado", "Saída stdout/stderr em tempo real"],
            ["Terminal do sistema", "Implementado", "Delegação a emulador interativo (TTY)"],
            ["Integração Git", "Implementado", "Status, pull, push, commit e discard"],
            ["Persistência SQLite", "Implementado", "settings, presets e run_log"],
            ["Workflow 8 etapas / GPM", "Fora do MVP", "Escopo original; não replicado"],
            ["Avaliação corporativa", "Não aplicável", "Escopo ajustado com orientador"],
        ],
        "table2": [
            [t["test_id"], t["name"], "Aprovado" if t["passed"] else "Reprovado", t["evidence"][:120]]
            for t in tests
        ],
        "table3": [
            [cmd, str(n)]
            for cmd, n in (metrics.get("comandos_mais_usados") or [])[:8]
        ]
        or [["(sem registros)", "0"]],
    }


def main() -> None:
    if not RESULTS_JSON.exists():
        raise SystemExit(f"Execute primeiro: tools/run_mvp_tests.py ({RESULTS_JSON} ausente)")
    if not TEMPLATE.exists():
        raise SystemExit(f"Template não encontrado: {TEMPLATE}")

    data = json.loads(RESULTS_JSON.read_text(encoding="utf-8"))
    content = build_content(data)

    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))

    setup_folha_de_rosto(doc)
    setup_second_page(doc)
    remove_instruction_paragraphs(doc)
    remove_optional_template_sections(doc)

    intro_heading = find_paragraph(doc, "Introdução")
    format_section_heading(intro_heading)
    anchor = insert_body_paragraphs_after(intro_heading, content["intro"])

    met_heading = find_paragraph(doc, "Metodologia ou Material e Métodos")
    format_section_heading(met_heading)
    anchor = insert_body_paragraphs_after(met_heading, content["metodologia_intro"])
    anchor = insert_subheading_after(anchor, "Desenvolvimento do protótipo")
    anchor = insert_body_paragraphs_after(anchor, content["metodologia_dev"])
    anchor = insert_subheading_after(anchor, "Validação técnica do protótipo")
    anchor = insert_body_paragraphs_after(anchor, content["metodologia_val"])

    res_heading = find_paragraph(doc, "Resultados Preliminares")
    format_section_heading(res_heading)
    anchor = insert_subheading_after(res_heading, "Desenvolvimento do protótipo")
    anchor = insert_body_paragraphs_after(anchor, content["resultados_dev"])
    anchor = insert_table_with_caption(
        anchor,
        "Tabela 1. Status de implementação dos componentes do MVP",
        ["Componente", "Status", "Observação"],
        content["table1"],
    )
    anchor = insert_subheading_after(anchor, "Validação técnica do protótipo")
    anchor = insert_body_paragraphs_after(anchor, content["resultados_val"][:2])
    anchor = insert_table_with_caption(
        anchor,
        "Tabela 2. Resultados do roteiro de testes funcionais (TF-01 a TF-09)",
        ["ID", "Caso de teste", "Resultado", "Evidência"],
        content["table2"],
    )
    anchor = insert_body_paragraphs_after(anchor, [content["resultados_val"][2]])
    anchor = insert_table_with_caption(
        anchor,
        "Tabela 3. Comandos mais frequentes no histórico run_log",
        ["Comando", "Ocorrências"],
        content["table3"],
        numeric_cols={1},
    )
    anchor = insert_body_paragraphs_after(anchor, [content["resultados_val"][3]])

    cons_heading = insert_paragraph_after(anchor, "Considerações Finais")
    format_section_heading(cons_heading)
    cons_body = insert_paragraph_after(cons_heading, CONSIDERACOES_FINAIS)
    format_body(cons_body)

    ref_heading = find_paragraph(doc, "Referências")
    format_section_heading(ref_heading)
    current = ref_heading
    for ref in content["refs"]:
        current = insert_paragraph_after(current, ref)
        format_reference(current)

    doc.save(str(OUTPUT))
    print(f"Gerado a partir do template: {OUTPUT}")


if __name__ == "__main__":
    main()
